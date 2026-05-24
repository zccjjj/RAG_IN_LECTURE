import os
import logging
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def format_time(seconds):
    """将秒数格式化为 MM:SS"""
    m = int(seconds) // 60
    s = int(seconds) % 60
    return f"{m:02d}:{s:02d}"


def generate_document(session_id, qa_results, output_dir="results"):
    """
    生成包含问答结果的 Word 文档。

    Args:
        session_id: 会话 ID
        qa_results: 问答结果列表 [{"question", "answer", "sources"}, ...]
        output_dir: 输出目录

    Returns:
        生成的文档路径
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # 标题
    title = doc.add_heading('视频问答结果', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加分隔线
    doc.add_paragraph('—' * 50)

    for i, qa in enumerate(qa_results, 1):
        # 问题标题
        q_heading = doc.add_heading(f'问题 {i}', level=1)

        # 问题内容
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(qa["question"])
        q_run.bold = True
        q_run.font.size = Pt(12)
        q_run.font.color.rgb = RGBColor(0, 51, 153)

        # 答案
        doc.add_heading('答案', level=2)
        a_para = doc.add_paragraph(qa["answer"])
        a_para.style.font.size = Pt(11)

        # 来源时间戳
        if qa.get("sources"):
            source_times = ", ".join(format_time(t) for t in qa["sources"])
            source_para = doc.add_paragraph()
            source_run = source_para.add_run(f"参考视频时间段: {source_times}")
            source_run.font.size = Pt(9)
            source_run.font.color.rgb = RGBColor(128, 128, 128)

        # 分隔
        if i < len(qa_results):
            doc.add_paragraph('—' * 50)

    # 保存文档
    output_path = os.path.join(output_dir, f"{session_id}.docx")
    doc.save(output_path)
    logger.info(f"Document saved to {output_path}")
    return output_path
